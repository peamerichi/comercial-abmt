from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

def add_table_row(table, cells_data, bold=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        if bold:
            run.bold = True
        if bg:
            shading = cell._element.get_or_add_tcPr()
            sh = shading.makeelement(qn('w:shd'), {qn('w:fill'): bg, qn('w:val'): 'clear'})
            shading.append(sh)
    return row

def set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(sh)

# ============ COVER ============
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DOCUMENTAÇÃO COMPLETA')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SISTEMA COMERCIAL ABMT')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Reconstrução do Projeto para Análise Externa')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Data: {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Empresa: A.B.M.T. Equipamentos Elétricos Ltda')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CNPJ: 17.820.873/0001-43')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============ TOC ============
doc.add_heading('ÍNDICE', level=1)
toc_items = [
    'PARTE 1 — Prompt Original / Intenção Original do Sistema',
    '  1.1 Objetivo Principal',
    '  1.2 Problemas que Deveria Resolver',
    '  1.3 Funcionalidades Solicitadas',
    '  1.4 Regras Comerciais',
    '  1.5 Fluxo de Vendas',
    '  1.6 Geração de Propostas',
    '  1.7 Controle de Comissão',
    '  1.8 Dados e Relatórios',
    '  1.9 Visão Estratégica',
    'PARTE 2 — Estado Atual do Sistema',
    '  2.1 Módulos Existentes',
    '  2.2 Telas Existentes',
    '  2.3 Funcionalidades Implementadas',
    '  2.4 Funcionalidades Parciais',
    '  2.5 Funcionalidades Planejadas',
    '  2.6 Fluxos Principais',
    '  2.7 Estrutura de Dados',
    '  2.8 Regras de Negócio',
    '  2.9 Integrações',
    '  2.10 Pontos Fortes e Fracos',
    '  2.11 Riscos',
    'PARTE 3 — Visão Funcional para Análise Externa',
    '  3A-3L: Seções descritivas completas',
    'TABELA MESTRE — Mapeamento Completo',
    'INVENTÁRIO — Arquivos, Componentes, Tabelas, Fluxos, Pendências',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('  '):
        p.runs[0].bold = True

doc.add_page_break()

# ============================================================
# PARTE 1
# ============================================================
doc.add_heading('PARTE 1 — PROMPT ORIGINAL / INTENÇÃO ORIGINAL DO SISTEMA', level=1)

p = doc.add_paragraph()
run = p.add_run('NOTA IMPORTANTE: ')
run.bold = True
run.font.color.rgb = RGBColor(0xE5, 0x3E, 0x3E)
p.add_run('O prompt original exato não foi preservado. O que segue é uma reconstrução fiel baseada na análise completa do código-fonte (app.py com ~5.300 linhas, app.js com ~5.100 linhas, database.py com ~700 linhas, forms.js, style.css), nas configurações seed do banco, nos nomes de variáveis/endpoints e na estrutura de dados. Tudo que é reconstrução está marcado como [RECONSTRUÍDO]. O que é verificável pelo código está marcado como [CONFIRMADO].')

doc.add_heading('1.1 Objetivo Principal', level=2)
doc.add_paragraph('[RECONSTRUÍDO] Criar um sistema web comercial completo (CRM + ERP comercial) para a ABMT, empresa do setor de equipamentos elétricos (transformadores, aço silício, cobre, alumínio, óleo isolante) com faturamento de ~R$80M/ano. O sistema deveria centralizar toda a operação comercial — desde a prospecção até o fechamento mensal — eliminando planilhas, WhatsApp como ferramenta de gestão, e controles manuais dispersos.')
doc.add_paragraph('[CONFIRMADO] A empresa é A.B.M.T. Equipamentos Elétricos Ltda, CNPJ 17.820.873/0001-43, localizada em Guarulhos/SP (Rua Veneza, 431). Estes dados estão hardcoded nas configurações seed do banco de dados.')

doc.add_heading('1.2 Problemas que Deveria Resolver', level=2)
problems = [
    'Falta de controle centralizado de propostas — vendedores gerenciavam negociações por WhatsApp e planilhas',
    'Sem rastreabilidade do ciclo de venda — não havia histórico de proposta → ordem → faturamento → recebimento',
    'Comissão calculada manualmente — cada fechamento exigia horas de trabalho manual com risco de erro',
    'Sem visibilidade de pipeline — gestor não sabia quantas propostas existiam, em que estágio, ou quanto esperar de faturamento',
    'Inadimplência invisível — parcelas vencidas não eram monitoradas sistematicamente',
    'Sem inteligência de recompra — clientes que paravam de comprar não eram detectados até ser tarde demais',
    'Dados dispersos — informações de clientes, preços, condições e histórico espalhados em múltiplas fontes',
    'Operação dependente de memória individual — se um vendedor saísse, todo o conhecimento ia embora',
]
for prob in problems:
    doc.add_paragraph(prob, style='List Bullet')
doc.add_paragraph('[RECONSTRUÍDO] Baseado nas funcionalidades implementadas e nas configurações do sistema.')

doc.add_heading('1.3 Funcionalidades Solicitadas no Início', level=2)
doc.add_paragraph('[RECONSTRUÍDO] Com base na análise do código, as funcionalidades originais incluíam:')
features = [
    ('Cadastro unificado de clientes/fornecedores', '[CONFIRMADO] Tabela cadastros serve ambos os papéis'),
    ('Propostas de venda e compra', '[CONFIRMADO] Tabela propostas com campo tipo=VENDA/COMPRA'),
    ('Conversão proposta → OV/OC', '[CONFIRMADO] Endpoint /api/propostas/<id>/converter'),
    ('Geração automática de parcelas', '[CONFIRMADO] Lógica de parcelamento por condição de pagamento'),
    ('Cálculo automático de comissão', '[CONFIRMADO] Tabela comissao_vendas com percentuais por categoria e perfil'),
    ('Fechamento mensal', '[CONFIRMADO] Tabela fechamentos + endpoint de fechamento com snapshot'),
    ('Dashboard com KPIs', '[CONFIRMADO] Múltiplos endpoints de dashboard'),
    ('Pipeline comercial', '[CONFIRMADO] Endpoint /api/pipeline com funil completo'),
    ('Controle de follow-ups', '[CONFIRMADO] Tabela followups com data/hora e conclusão'),
    ('Busca global', '[CONFIRMADO] FTS5 para busca full-text em todo o sistema'),
    ('Notas pessoais', '[CONFIRMADO] Tabela notas com tags, cores, checklist, vínculos'),
    ('Sistema de notificações', '[CONFIRMADO] Tabela notificacoes com lida/não-lida'),
    ('Configurações editáveis', '[CONFIRMADO] Tabela configuracoes com ICMS, PIS, comissão, política'),
]
for feat, status in features:
    p = doc.add_paragraph()
    run = p.add_run(f'{feat}: ')
    run.bold = True
    p.add_run(status)

doc.add_heading('1.4 Regras Comerciais Consideradas', level=2)
doc.add_paragraph('[CONFIRMADO] As seguintes regras comerciais estão implementadas no código:')

rules = [
    ('ICMS interestadual', 'Tabela completa de alíquotas por UF de destino (SP=18%, Sul/Sudeste=12%, restante=7%). Toggle SP Normal/Isento. Origem sempre SP.'),
    ('PIS/COFINS', 'Percentual configurável, padrão 9.25%. Aplicado sobre valor bruto para cálculo de valor líquido.'),
    ('Comissão por perfil e categoria', '3 perfis (vendedor, gerente, diretor) com percentuais diferentes por categoria de produto. Comissão calculada sobre valor líquido (após PIS e ICMS). Ex: Transformador Usado = 3% gerente, 2.5% vendedor.'),
    ('Comissão de compras', 'Percentuais fixos por nome: Pedro=3%, Thiago=2.5%, Guilherme=3%, diferença gestor=0.5%.'),
    ('Condições de pagamento', 'À vista, 30d, 30/60d, 30/60/90d, 30/60/90/120d, 28/56d, personalizado. Parcelas geradas automaticamente na conversão.'),
    ('Formas de pagamento', 'Faturado, PIX, Boleto, Depósito, TED, Cartão, Outro.'),
    ('Frete', 'FOB ou CIF, com valor e transportadora opcionais.'),
    ('Validade de proposta', 'Padrão 7 dias. Expiração automática ao vencer.'),
    ('Limite de crédito', 'Campo limite_faturamento por cadastro. Sistema calcula crédito tomado (parcelas pendentes), liberado e disponível.'),
    ('Scoring qualitativo', 'Fórmula H = 10 × entrada% + base[faturamento] × (1 - entrada%). Base varia de 7.0 (à vista) a 0.0 (75d).'),
    ('Bônus semanal', 'Meta semanal de R$50.000 com bônus de R$250.'),
]
for title, desc in rules:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('1.5 Fluxo de Vendas', level=2)
doc.add_paragraph('[CONFIRMADO] O fluxo completo implementado no sistema:')
steps = [
    '1. Cadastrar cliente (CNPJ/CPF, razão social, endereço, contato, segmento)',
    '2. Criar proposta de venda (selecionar cliente, adicionar itens com categoria/quantidade/preço, definir condição de pagamento, frete, UF destino)',
    '3. Workflow de proposta: Rascunho → Enviada → Em Negociação → Aprovada → Convertida (ou Perdida/Expirada)',
    '4. Converter proposta aprovada em Ordem de Venda (OV) — cálculo automático de comissão e geração de parcelas',
    '5. Workflow de OV: Aprovada → Em Produção → Faturada → Despachada → Entregue (ou Cancelada)',
    '6. Acompanhar parcelas (contas a receber) com baixa total ou parcial',
    '7. Fechamento mensal com snapshot de comissões e exportação CSV',
]
for step in steps:
    doc.add_paragraph(step)

doc.add_heading('1.6 Geração de Propostas', level=2)
doc.add_paragraph('[CONFIRMADO] O sistema gera propostas com os seguintes elementos:')
prop_items = [
    'Numeração sequencial automática (PROP-0001, PROP-0002...)',
    '13 categorias de produto: Transformador Usado/Novo, Bobinas/Chapas de Aço Silício, Chapas Cortadas, Caixa e Núcleo, Cobre, Alumínio, Óleo Isolante, Radiadores, Papel Kraft, Retalho/Sucata, Diversos',
    'Campos específicos por categoria armazenados em JSON (campos_especificos)',
    'Unidades: KG, KVA, LITRO, UNIDADE',
    'Desconto por item (percentual ou valor fixo)',
    'Cálculo automático de peso total e valor total',
    'Custo de embalagem para Óleo Isolante (adicionado ao valor total)',
    'Obs para cliente (aparece em output) e obs interna (não aparece)',
    'Opção de incluir dados bancários e política comercial',
    'Vinculação entre proposta de compra e venda (proposta_vinculada_id)',
    'Intermediário com comissão (para vendas com brokers)',
    'Duplicação de proposta em 1 clique',
    'Revalidação de proposta expirada',
    'Conversão direta em OV ou OC com geração de parcelas',
]
for item in prop_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.7 Controle de Comissão', level=2)
doc.add_paragraph('[CONFIRMADO] O sistema implementa comissão automática:')
com_details = [
    'Tabela de comissão configurável por perfil (vendedor/gerente/diretor) e por categoria de produto',
    'Base de cálculo: valor líquido do item (após PIS/COFINS e ICMS)',
    'Fórmula: comissão = valor_líquido_item × percentual_categoria / 100',
    'Percentuais seed: Transformador Usado (gerente 3%, vendedor 2.5%), Cobre/Alumínio/Óleo (gerente 0.5%, vendedor 0.4%)',
    'Comissão calculada automaticamente na conversão proposta→OV e na criação direta de OV',
    'Fechamento mensal gera snapshot JSON com todas as comissões do período',
    'Dados sensíveis (custo, margem, comissão) ocultos do perfil vendedor via API',
]
for item in com_details:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.8 Dados e Relatórios', level=2)
doc.add_paragraph('[CONFIRMADO] Relatórios e analytics implementados:')
reports = [
    'Dashboard em tempo real: vendas/compras do mês, follow-ups, parcelas vencidas, notificações',
    'Dashboard avançado: filtros por vendedor, categoria, segmento, mês, ano',
    'Analytics anual: vendas/compras por mês, por categoria, por vendedor, top 10 clientes/fornecedores, ticket médio, taxa conversão',
    'Analytics trimestral: mesmas métricas segmentadas por trimestre, comparativo com trimestre anterior',
    'Dashboard comparativo: período atual vs anterior vs mesmo mês ano anterior',
    'Pipeline comercial: propostas por período (hoje/semana/mês), faturamento real, por vendedor, abertas, crédito de clientes, funil por status',
    'Previsão de faturamento: propostas abertas hoje/semana/pipeline total',
    'Fluxo de caixa: projeção 90 dias com entradas (OV parcelas) vs saídas (OC parcelas), gráfico semanal',
    'CRM Alertas: recompra (clientes fora do ciclo), inativos (90+ dias), vencimentos próximos (7 dias)',
    'Fechamento mensal: comissões de vendas e compras, scoring qualitativo, exportação CSV',
    'Consulta rápida (IA Insights): respostas por keyword matching sobre faturamento, clientes, pipeline, comissão, metas',
]
for item in reports:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.9 Visão Estratégica', level=2)
doc.add_paragraph('[RECONSTRUÍDO] A visão estratégica por trás do sistema era transformar a operação comercial da ABMT de um modelo baseado em memória individual e comunicação informal (WhatsApp) para um modelo baseado em processo e dados. Os principais pilares estratégicos:')
strategic = [
    'Centralização: toda a operação comercial em um único sistema acessível de qualquer dispositivo (PWA)',
    'Rastreabilidade: cada real movimentado pode ser rastreado desde a proposta até o recebimento',
    'Automação: comissão, parcelas, scoring e alertas calculados automaticamente',
    'Visibilidade: gestor/diretor tem visão em tempo real de toda a operação sem precisar perguntar',
    'Inteligência: sistema detecta clientes em risco, inadimplência e oportunidades de recompra',
    'Escalabilidade: suportar crescimento da equipe comercial sem perder controle',
    'Autonomia: vendedor opera independente, gestor monitora remotamente',
]
for item in strategic:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# PARTE 2
# ============================================================
doc.add_heading('PARTE 2 — ESTADO ATUAL DO SISTEMA', level=1)

doc.add_heading('2.1 Módulos Existentes', level=2)
modules = [
    ('Autenticação', 'Login com sessão, 3 perfis (vendedor/gerente/diretor), CSRF protection, troca de senha obrigatória no primeiro acesso'),
    ('Cadastros', 'CRUD de clientes/fornecedores unificado, validação CNPJ/CPF com dígito verificador, LTV health tracking, dual role (cliente/fornecedor/ambos), contatos adicionais em JSON'),
    ('Propostas', 'CRUD completo de propostas de venda e compra, workflow de status com transições validadas, itens com campos específicos por categoria, duplicação, vinculação compra/venda, intermediário'),
    ('Ordens de Venda', 'Criadas via conversão de proposta ou diretamente, workflow expandido (6 status), edição de campos, comissão automática, NF, número Omie, parcelas com baixa parcial'),
    ('Ordens de Compra', 'Mesma estrutura de OV para compras, recebimento parcial de itens, intermediário, NF, parcelas a pagar'),
    ('Financeiro', 'Contas a receber (OV parcelas) e contas a pagar (OC parcelas), marcação batch, baixa parcial, auto-update de vencidas, cashflow 90 dias'),
    ('Dashboard', 'Dashboard básico, avançado com filtros (vendedor/categoria/segmento/mês/ano), comparativo, margem bruta, ranking vendedores, inadimplência, ações rápidas'),
    ('Pipeline', 'Visão completa de funil: propostas por período, faturamento real, performance por vendedor, pipeline ativo, análise de crédito'),
    ('Analytics', 'Relatórios anual e trimestral, vendas/compras por mês/categoria/vendedor/UF, top clientes/fornecedores, taxa conversão, ticket médio'),
    ('CRM/Inteligência', 'Alertas de recompra baseados em frequência, clientes inativos, vencimentos próximos, repetir pedido automático, interações (timeline CRM), consulta rápida (keyword matching)'),
    ('Follow-ups', 'Lista de tarefas por vendedor com data/hora, vinculação a cadastro/proposta/OV/OC, conclusão com registro'),
    ('Notas', 'Notas pessoais com título, conteúdo, tags JSON, cores, checklist, fixação, vínculo a entidades, lembrete'),
    ('Fechamento', 'Fechamento mensal de vendas/compras, cálculo de comissão por vendedor, scoring qualitativo, snapshot JSON, export CSV, bloqueio/reabertura'),
    ('Configurações', 'ICMS por UF, PIS/COFINS, comissão por perfil/categoria, dados empresa, dados bancários, política comercial, templates de condições'),
    ('Busca Global', 'FTS5 full-text search em cadastros, propostas, OVs, OCs, com fallback para LIKE'),
    ('Notificações', 'Sistema de notificações in-app com badge, dropdown, marcar como lida'),
    ('Sugestões', 'Canal de feedback interno (Bug, Melhoria, Dúvida, Outro) com status de análise'),
    ('Anexos', 'Upload de arquivos vinculados a OV/OC/proposta, soft delete, limite 5MB'),
    ('Backup', 'Backup automático a cada 12h, rotação de 30 cópias, backup on start se >6h'),
]
for name, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 Telas Existentes', level=2)
screens = [
    ('Login', 'Tela de login com "lembrar usuário", frases motivacionais, troca de senha obrigatória'),
    ('Dashboard', 'Tela principal com saudação, filtros, tabs vendas/compras, KPIs, gráficos Chart.js, ranking, inadimplência, ações rápidas, previsão, cashflow'),
    ('Propostas de Venda', 'Lista paginada com busca/filtro, formulário completo, visualização detalhada com log, taxas, crédito do cliente'),
    ('Propostas de Compra', 'Mesma estrutura de vendas para compras'),
    ('Ordens de Venda', 'Lista paginada, visualização com status workflow, parcelas, baixa parcial, edição, OCs vinculadas'),
    ('Ordens de Compra', 'Lista paginada, visualização com recebimento, NF, edição'),
    ('Clientes/Cadastros', 'Lista com busca, filtro por papel/segmento/status, LTV health badges, vista individual com timeline CRM, crédito, histórico OVs/OCs'),
    ('Follow-ups', 'Lista de tarefas pendentes, separação hoje/atrasados/futuros, conclusão'),
    ('Notas', 'Grid de notas pessoais estilo post-it, cores, fixação, checklist'),
    ('Pipeline Comercial', 'Visão de funil, KPIs diários/semanais/mensais, crédito por cliente (gestor only)'),
    ('Fechamento', 'Seletor mês/ano, tabela de comissões, scoring, export CSV (gestor only)'),
    ('Relatórios', 'Analytics anual/trimestral com gráficos e tabelas (gestor only)'),
    ('Inteligência Comercial', 'Alertas CRM, clientes inativos, vencimentos (gestor only)'),
    ('Configurações', 'Edição de ICMS, PIS, comissão, empresa, banco, política (gestor only)'),
    ('Guia do Vendedor', 'Página estática com orientações de uso do sistema'),
    ('Notificações', 'Lista de notificações com marcar lida'),
    ('Assistente/Sugestões', 'Panel flutuante com chat de consulta rápida e formulário de sugestões'),
]
for name, desc in screens:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.3 Funcionalidades Implementadas', level=2)
implemented = [
    'CRUD completo de cadastros com validação CNPJ/CPF e contatos adicionais',
    'CRUD completo de propostas com itens, campos específicos por categoria, workflow de status com transições validadas',
    'Conversão atômica proposta→OV/OC com race condition prevention',
    'Duplicação de propostas',
    'Vinculação proposta compra↔venda',
    'Intermediário com comissão em propostas',
    'Criação direta de OV/OC (sem proposta)',
    'Workflow expandido de OV: Aprovada→Em Produção→Faturada→Despachada→Entregue→Cancelada',
    'Edição de OV e OC após criação',
    'Parcelas a receber e a pagar com auto-vencimento',
    'Baixa total e parcial de parcelas com audit log',
    'Marcação batch de parcelas',
    'Cashflow 90 dias com gráfico (Chart.js)',
    'Dashboard com filtros combinados (vendedor, categoria, segmento, mês, ano)',
    'Ranking de vendedores com meta e progresso visual',
    'Margem bruta estimada no dashboard',
    'Inadimplência visível no dashboard do gestor',
    'Compartilhar resumo para WhatsApp (copiar texto formatado)',
    'Alertas de recompra baseados em frequência de compra',
    'Repetir pedido em 1 clique (gera proposta baseada na última OV)',
    'Detecção de clientes inativos (90+ dias)',
    'Análise de crédito por cliente (tomado/liberado/disponível)',
    'LTV health tracking com badges (ok/atenção/risco)',
    'Interações CRM (ligação, WhatsApp, email, reunião, visita)',
    'Follow-ups com vinculação a entidades',
    'Notas pessoais com checklist e cores',
    'Busca global FTS5 + fallback LIKE',
    'Notificações in-app',
    'Expiração automática de propostas',
    'Scoring qualitativo de condição de pagamento',
    'Comissão automática por perfil e categoria',
    'Fechamento mensal com snapshot e CSV',
    'Analytics anual e trimestral com comparativos',
    'Consulta rápida (keyword matching) para dados comerciais',
    'Tema claro/escuro toggle',
    'Pull-to-refresh no mobile',
    'Service Worker com cache e update banner',
    'CSRF protection em todas as rotas POST/PUT/DELETE',
    'Proteção double-click com safeAction()',
    'Formulário dirty check (avisa antes de sair sem salvar)',
    'Configurações editáveis (ICMS, PIS, comissão, empresa, banco, política)',
    'Canal de sugestões/feedback interno',
    'Backup automático com rotação',
    'OC-OV Links para tracking de margem',
]
for item in implemented:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.4 Funcionalidades Parcialmente Implementadas', level=2)
partial = [
    ('Metas por vendedor', 'Tabela metas existe no banco. Dashboard mostra ranking com meta. Porém não existe endpoint dedicado de CRUD de metas — são criadas via configurações genéricas. Status: PARCIAL'),
    ('OC-OV Links', 'Tabela e endpoint existem, mas a UX de vinculação é básica. Não há sugestão automática de quais OCs vincular. Status: PARCIAL'),
    ('Recebimento parcial de OC', 'Endpoint existe para registrar quantidade recebida, mas a UX de controle item a item é limitada. Status: PARCIAL'),
    ('Busca global FTS5', 'Funciona, mas a indexação nem sempre é consistente (depende de update_fts ser chamado). Fallback LIKE compensa. Status: PARCIAL'),
    ('Expiração de propostas', 'Auto-expira ao listar/visualizar, mas não roda como job independente (depende de alguém acessar). Status: PARCIAL'),
]
for name, desc in partial:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.5 Funcionalidades Planejadas mas Não Implementadas', level=2)
planned = [
    ('Tela "Meu Dia" para vendedor', 'Dashboard focado em ações concretas do dia. Status: PLANEJADO'),
    ('Follow-up automático por status de proposta', 'Criar follow-up ao mudar status. Status: PLANEJADO'),
    ('Pipeline ponderado com probabilidade', 'Pesar pipeline por chance de fechamento por status. Status: PLANEJADO'),
    ('DRE Comercial', 'Demonstrativo de resultado: receita - custo - comissão - impostos. Status: PLANEJADO'),
    ('Aging de contas a receber', 'Buckets 1-30, 31-60, 61-90, 90+ dias. Status: PLANEJADO'),
    ('Motivo de perda estruturado', 'Dropdown em vez de texto livre. Status: PLANEJADO'),
    ('Relatório ad-hoc com filtros combinados', 'Explorador de dados livre. Status: PLANEJADO'),
    ('Push notifications reais', 'Web Push API via Service Worker. Status: PLANEJADO'),
    ('Integração Omie', 'Sincronização de NF e parcelas. Status: PLANEJADO'),
    ('Multi-empresa (ABMT + AEB)', 'Tenant isolation. Status: PLANEJADO'),
    ('API documentada (Swagger)', 'OpenAPI/Swagger. Status: PLANEJADO'),
    ('Migração para PostgreSQL', 'Substituir SQLite. Status: PLANEJADO'),
]
for name, desc in planned:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.6 Fluxos Principais do Sistema', level=2)

doc.add_paragraph('Fluxo 1 — Venda completa:', style='List Bullet')
doc.add_paragraph('Cadastrar cliente → Criar proposta VENDA (itens, condição, frete) → Enviar (Rascunho→Enviada) → Negociar (→Em Negociação) → Aprovar (→Aprovada) → Converter em OV (parcelas e comissão automáticas) → Avançar status OV (Produção→Faturada→Despachada→Entregue) → Acompanhar parcelas → Baixar pagamentos → Fechamento mensal')

doc.add_paragraph('Fluxo 2 — Compra completa:', style='List Bullet')
doc.add_paragraph('Cadastrar fornecedor → Criar proposta COMPRA → Converter em OC → Receber material (parcial ou total) → Acompanhar contas a pagar → Baixar pagamentos')

doc.add_paragraph('Fluxo 3 — Recompra:', style='List Bullet')
doc.add_paragraph('Sistema detecta cliente fora do ciclo de compra → Alerta no dashboard → Vendedor clica "Repetir Pedido" → Proposta criada automaticamente com itens da última OV → Vendedor ajusta e envia')

doc.add_paragraph('Fluxo 4 — Fechamento mensal:', style='List Bullet')
doc.add_paragraph('Gestor acessa Fechamento → Seleciona mês/ano → Sistema calcula comissões de vendas (por vendedor, por OV, por item, com scoring qualitativo) e compras → Gestor revisa → Fecha o mês (gera snapshot JSON) → Exporta CSV para pagamento')

doc.add_heading('2.7 Estrutura de Dados', level=2)
doc.add_paragraph('O banco SQLite (comercial.db) contém 20+ tabelas:')
tables = [
    ('users', '3 perfis: vendedor, gerente, diretor. Campos: username, password_hash, nome, perfil, cpf, dados_bancarios, ativo, must_change_password'),
    ('cadastros', 'Clientes e fornecedores unificados. CNPJ/CPF único, tipo_pessoa PF/PJ, endereço completo, contato principal + adicionais (JSON), segmento, tags (JSON), vendedor responsável, limite faturamento, status Ativo/Inativo/Bloqueado'),
    ('propostas', 'Número sequencial, tipo COMPRA/VENDA, status com CHECK (7 valores), vinculada a cadastro e vendedor, condição pagamento JSON, intermediário, validade com expiração'),
    ('proposta_items', 'Vinculada a proposta, categoria, campos_especificos JSON, peso, quantidade, unidade, valor com desconto'),
    ('proposta_log', 'Histórico de ações na proposta (criação, edição, conversão, status)'),
    ('ordens_venda', 'Status expandido (6 valores incluindo Em Produção, Faturada, Despachada), nota_fiscal, numero_omie, cancelamento com motivo e responsável'),
    ('ov_items', 'Igual proposta_items + custo, margem, comissao_percentual, comissao_valor, oc_origem_id, status por item'),
    ('ov_parcelas', 'Contas a receber. Status: Pendente/Vencida/Paga/Paga Parcial. valor_recebido para baixa parcial'),
    ('ordens_compra', 'Espelho de OV para compras. comprador_id, intermediário com comissão, nota_fiscal'),
    ('oc_items', 'Igual ov_items + quantidade_recebida, status de recebimento'),
    ('oc_parcelas', 'Contas a pagar. Mesma estrutura de ov_parcelas com valor_pago'),
    ('oc_ov_links', 'Vínculo OC↔OV para tracking de margem. valor_alocado_compra/venda. UNIQUE(oc_id, ov_id)'),
    ('anexos', 'Arquivos vinculados a OV/OC/proposta. Soft delete'),
    ('followups', 'Tarefas agendadas por vendedor. Vínculo a cadastro/proposta/OV/OC'),
    ('interacoes', 'Timeline CRM: Ligação, WhatsApp, Email, Reunião, Visita, Outro'),
    ('notas', 'Notas pessoais com tags JSON, checklist JSON, cor, fixação, lembrete'),
    ('fechamentos', 'Fechamento mensal por tipo (geral/vendas/compras). Snapshot JSON das comissões'),
    ('configuracoes', 'Key-value para todas as configs do sistema'),
    ('audit_log', 'Log de auditoria por entidade'),
    ('notificacoes', 'Notificações in-app por usuário'),
    ('busca_global', 'FTS5 virtual table para busca full-text'),
    ('metas', 'Meta mensal/semanal por vendedor'),
    ('sugestoes', 'Feedback de usuários'),
    ('condicoes_salvas', 'Condições de pagamento favoritas'),
]
for name, desc in tables:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    run.font.size = Pt(9)
    r2 = p.add_run(desc)
    r2.font.size = Pt(9)

doc.add_heading('2.8 Regras de Negócio Existentes', level=2)
biz_rules = [
    'Vendedor só vê seus próprios dados (propostas, OVs); gestor/diretor vê tudo',
    'Vendedor não vê custo, margem, comissão nos itens da OV',
    'Transições de status de proposta são validadas (ex: Rascunho só pode ir para Enviada)',
    'Gestor pode reverter proposta de Aprovada para Em Negociação',
    'Proposta Convertida ou Perdida não pode ser editada',
    'Conversão de proposta é atômica (race condition prevention)',
    'Parcelas vencidas são marcadas automaticamente ao listar (UPDATE antes do SELECT)',
    'Comissão é sobre valor líquido (após PIS e ICMS)',
    'Fechamento mensal pode ser reaberto por gestor',
    'Backup automático a cada 12h com rotação de 30 cópias',
    'CNPJ/CPF validado com algoritmo de dígito verificador',
    'Upload limitado a 5MB, extensões permitidas: png, jpg, jpeg, gif, pdf, doc, docx, xls, xlsx',
    'Sessão expira em 8 horas',
    'CSRF token obrigatório em todas as operações de escrita',
]
for item in biz_rules:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.9 Integrações Existentes ou Previstas', level=2)
integrations = [
    ('WhatsApp', 'IMPLEMENTADO — link wa.me/ direto para o contato do cliente/fornecedor a partir de diversas telas (recompra, cadastro view)'),
    ('Chart.js', 'IMPLEMENTADO — gráficos de barras e linhas no dashboard, cashflow, analytics. Carregado via CDN.'),
    ('Omie', 'PREVISTO — campos numero_omie já existem em OV e OC, mas não há integração real com a API Omie'),
    ('Email', 'NÃO EXISTE — sem integração de email para envio de propostas ou relatórios'),
    ('ERP/NF', 'NÃO EXISTE — campo nota_fiscal é texto livre manual'),
    ('BI Tools', 'NÃO EXISTE — sem API documentada para conexão externa'),
]
for name, desc in integrations:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.10 Pontos Fortes e Fracos Atuais', level=2)
doc.add_paragraph('PONTOS FORTES:', style='Heading 3')
strengths = [
    'Ciclo completo proposta→ordem→parcela→fechamento implementado',
    'Comissão automática por perfil e categoria com scoring qualitativo',
    'Análise de crédito por cliente integrada ao pipeline',
    'Alertas de recompra baseados em frequência real',
    'Baixa parcial de parcelas com audit log',
    'Performance otimizada: bulk prefetch, GROUP BY em vez de N+1 queries',
    'CSRF + session + role-based access + safeAction()',
    'PWA com Service Worker e update banner',
    'FTS5 para busca global',
    'Dashboard rico com filtros combinados e compartilhamento WhatsApp',
]
for item in strengths:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('PONTOS FRACOS:', style='Heading 3')
weaknesses = [
    'SQLite em produção — write lock global com múltiplos usuários simultâneos',
    'Monolito: app.py com 5.300 linhas sem blueprints',
    'app.js com 5.100 linhas sem componentização ou state management',
    'Zero testes automatizados',
    'Meta por vendedor parcialmente implementada (tabela existe, CRUD incompleto)',
    'Motivo de perda é texto livre (não estruturado)',
    'Pipeline não ponderado (sem probabilidade por status)',
    'Sem DRE comercial (receita - custo - comissão - impostos)',
    'Sem aging de contas a receber',
    'Follow-up manual (não automático por workflow)',
    'IA Insights é keyword matching (nome "IA" gera expectativa incorreta)',
    'Categorias de produto hardcoded no Python',
    'Sem push notifications reais',
    'Sem integração real com ERP/NF/email',
]
for item in weaknesses:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.11 Riscos', level=2)
doc.add_paragraph('RISCOS TÉCNICOS:', style='Heading 3')
tech_risks = [
    'SQLite lock contention com 5+ usuários simultâneos em horário de pico',
    'Monolito de 5300 linhas dificulta manutenção e introduz acoplamento',
    'Sem testes = qualquer mudança pode quebrar comissão ou parcelas sem detecção',
    'Service Worker pode cachear versão antiga se versionamento não for atualizado',
    'Backup interno sem recovery testado',
]
for item in tech_risks:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('RISCOS OPERACIONAIS:', style='Heading 3')
op_risks = [
    'Adoção: vendedor pode não abrir o sistema diariamente sem valor imediato',
    'Dados incompletos: cadastros sem segmento/UF tornam análises inúteis',
    'Dependência do Pedro: faturamento depende de ação comercial diária dele',
    'Pipeline inflado: propostas expiradas que ainda não foram acessadas não mudam status',
    'Motivo de perda como texto livre impossibilita análise "por que perdemos"',
]
for item in op_risks:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# PARTE 3
# ============================================================
doc.add_heading('PARTE 3 — VISÃO FUNCIONAL PARA ANÁLISE EXTERNA', level=1)

doc.add_heading('A. O que é o sistema', level=2)
doc.add_paragraph('O Sistema Comercial ABMT é uma aplicação web (PWA — Progressive Web App) que funciona como CRM e ERP comercial integrado para a empresa A.B.M.T. Equipamentos Elétricos Ltda. Foi construído como uma SPA (Single Page Application) com backend Flask (Python) e banco SQLite, acessível via navegador em qualquer dispositivo (desktop, tablet, celular). O sistema cobre o ciclo comercial completo: prospecção e cadastro de clientes/fornecedores, criação e gestão de propostas, conversão em ordens de venda/compra, controle financeiro (parcelas a receber e pagar), cálculo de comissão, fechamento mensal e inteligência comercial.')
doc.add_paragraph('Tecnologias: Flask 3.x (Python), SQLite com WAL mode e FTS5, JavaScript vanilla (SPA), CSS custom, Chart.js para gráficos, Service Worker para cache, Lucide Icons (SVG inline).')

doc.add_heading('B. Para quem ele foi criado', level=2)
doc.add_paragraph('O sistema atende 3 perfis de usuário na ABMT:')
users = [
    ('Vendedor (ex: Thiago)', 'Cria propostas, acompanha seus pedidos, registra follow-ups, consulta dados dos seus clientes. Não vê custo, margem ou comissão detalhada. Não acessa pipeline, fechamento ou configurações.'),
    ('Gerente (ex: Pedro)', 'Tudo que o vendedor faz + acessa pipeline, fechamento, relatórios, inteligência comercial, configurações. Vê custos, margens e comissões. Pode aprovar/reverter propostas de outros vendedores.'),
    ('Diretor (ex: Guilherme)', 'Mesmo acesso do gerente. Perfil CEO com visão total do sistema.'),
]
for name, desc in users:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)
doc.add_paragraph('A ABMT opera no setor de equipamentos elétricos, comprando e vendendo transformadores (usados e novos), aço silício (bobinas e chapas), cobre, alumínio, óleo isolante e materiais correlatos. O público-alvo são reformadores de transformadores, fabricantes, concessionárias de energia, indústrias e recicladores.')

doc.add_heading('C. Como a venda acontece na prática', level=2)
doc.add_paragraph('Na prática, a venda na ABMT segue este fluxo:')
sale_flow = [
    '1. O vendedor identifica uma oportunidade (contato do cliente, lead existente, recompra sugerida pelo sistema)',
    '2. Negocia via WhatsApp/telefone os termos: qual material, quantidade, preço por kg/kVA/litro/unidade, condição de pagamento, frete',
    '3. Quando há acordo preliminar, o vendedor cria uma proposta no sistema com todos os itens, preços e condições',
    '4. A proposta avança no workflow: Rascunho → Enviada (ao cliente) → Em Negociação (se houver contrapartida) → Aprovada (se fechou)',
    '5. A proposta aprovada é convertida em Ordem de Venda (OV), que gera automaticamente as parcelas financeiras e calcula a comissão',
    '6. A OV avança: Aprovada → Em Produção (se precisa preparar) → Faturada (NF emitida) → Despachada (material enviado) → Entregue',
    '7. O financeiro acompanha as parcelas e registra recebimentos (baixa total ou parcial)',
    '8. No final do mês, o gestor faz o fechamento que consolida comissões e gera relatório',
]
for step in sale_flow:
    doc.add_paragraph(step)

doc.add_heading('D. Como o WhatsApp entra no processo', level=2)
doc.add_paragraph('O WhatsApp é a principal ferramenta de comunicação com clientes na ABMT. O sistema integra com WhatsApp de forma leve:')
wa_items = [
    'Cada cadastro tem campo contato_whatsapp',
    'Em várias telas (recompra, cadastro view), há botão direto para abrir conversa via wa.me/55{numero}',
    'O dashboard tem função "Compartilhar resumo" que copia texto formatado para colar no WhatsApp do time',
    'Interações do tipo "WhatsApp" podem ser registradas na timeline CRM do cliente',
    'NÃO EXISTE integração automática (envio de proposta por WhatsApp, tracking de mensagens, chatbot)',
]
for item in wa_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('E. Como o vendedor formaliza a negociação', level=2)
doc.add_paragraph('Após negociar via WhatsApp/telefone, o vendedor formaliza criando uma proposta no sistema. O formulário de proposta exige: seleção do cliente (por CNPJ/CPF com busca automática), adição de itens com categoria específica (ex: Transformador Usado, Bobinas de Aço Silício), quantidade, unidade (KG/KVA/LITRO/UNIDADE), valor unitário, desconto opcional. Também define: condição de pagamento (à vista ou parcelado), forma de pagamento (faturado/PIX/boleto/etc), frete (FOB/CIF), UF de destino (para cálculo de ICMS), validade da proposta, observações para o cliente e observações internas.')
doc.add_paragraph('O sistema calcula automaticamente: valor total por item, peso total, impostos estimados (PIS/COFINS + ICMS), valor líquido, e comissão estimada (visível apenas para gestor).')

doc.add_heading('F. Como a proposta é criada', level=2)
doc.add_paragraph('A proposta é criada através de um formulário multi-seção no frontend (renderPropostaForm em forms.js):')
prop_sections = [
    'Seção 1 — Cliente/Fornecedor: busca por CNPJ/CPF com lookup automático, mostra resumo do cadastro',
    'Seção 2 — Responsável: seleção do vendedor, UF destino (com toggle SP Normal/Isento para ICMS)',
    'Seção 3 — Itens: acordeão com campos dinâmicos por categoria. Ex: Transformador tem campos de potência, tensão, tipo; Óleo tem campo de embalagem com custo',
    'Seção 4 — Condições: tipo de pagamento (dropdown com condições rápidas ou personalizado), forma, dados bancários/PIX, frete, transportadora',
    'Seção 5 — Observações: obs para cliente (aparece no output), obs interna, validade, garantia, prazo entrega',
    'Seção 6 — Intermediário (opcional): cadastro intermediário, comissão, obs',
    'Seção 7 — Opções: incluir dados bancários, incluir política comercial, mostrar impostos',
]
for item in prop_sections:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('O backend gera número sequencial (PROP-0001), salva com todos os itens, registra no log, atualiza índice FTS5.')

doc.add_heading('G. Como o gestor acompanha a operação', level=2)
doc.add_paragraph('O gestor tem acesso a múltiplas ferramentas de monitoramento:')
gestor_tools = [
    'Dashboard com filtros: vê vendas e compras do mês por vendedor/categoria/segmento, com comparativo',
    'Ranking de vendedores: tabela com meta vs realizado e barra de progresso visual',
    'Inadimplência: lista de parcelas vencidas com dias de atraso, direto no dashboard',
    'Margem bruta estimada: faturamento - compras no mês',
    'Pipeline comercial: funil de propostas abertas, faturamento real por período, análise de crédito por cliente',
    'Previsão de faturamento: propostas que podem virar venda hoje/semana/pipeline total',
    'Fluxo de caixa 90 dias: gráfico de entradas vs saídas projetadas por semana',
    'Alertas de recompra: clientes fora do ciclo habitual de compra',
    'Clientes inativos: quem não compra há 90+ dias',
    'Relatórios: analytics anual e trimestral com gráficos, comparativos, top clientes/fornecedores',
    'Fechamento mensal: comissões calculadas, scoring qualitativo, export CSV',
]
for item in gestor_tools:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('H. Como os dados comerciais são gerados', level=2)
doc.add_paragraph('Os dados são gerados a partir das operações do dia a dia:')
data_gen = [
    'Cada proposta criada/editada/convertida gera registro em proposta_log',
    'Cada OV gera itens com comissão calculada e parcelas financeiras',
    'Cada baixa de parcela gera registro em audit_log',
    'O dashboard agrega dados em tempo real via queries SQL com GROUP BY',
    'Analytics usa as mesmas tabelas com filtros de período (strftime para SQLite)',
    'Fechamento gera snapshot JSON imutável das comissões do mês',
    'Não existe ETL ou data warehouse — tudo é query direto no banco operacional',
]
for item in data_gen:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('I. Como a comissão é calculada', level=2)
doc.add_paragraph('O cálculo de comissão segue esta lógica implementada no código:')
com_logic = [
    '1. Identificar o perfil do vendedor (vendedor/gerente/diretor)',
    '2. Para cada item da OV, buscar o percentual de comissão na tabela configurável (por perfil × categoria)',
    '3. Calcular valor líquido do item: valor_total × (1 - PIS%/100 - ICMS%/100)',
    '4. Comissão do item = valor_líquido × percentual / 100',
    '5. Comissão total da OV = soma das comissões dos itens',
    '6. No fechamento mensal, todas as OVs do mês são consolidadas por vendedor',
    '7. O scoring qualitativo avalia a qualidade da condição de pagamento (à vista = nota 10, 75d = nota 0)',
    '8. Existe também comissão de compras com percentuais fixos por nome de comprador',
]
for item in com_logic:
    doc.add_paragraph(item)

doc.add_heading('J. Como o sistema apoia operação baseada em LTV', level=2)
doc.add_paragraph('O sistema implementa inteligência de LTV (Lifetime Value) de várias formas:')
ltv_items = [
    'Cada cadastro na lista mostra: dias sem comprar, total de pedidos, LTV total (soma de todas as OVs), frequência média de compra',
    'Health badge: "ok" (dentro do ciclo), "atenção" (10+ dias além), "risco" (30+ dias além do ciclo médio)',
    'Alertas de recompra: calcula intervalo médio entre compras e alerta quando o cliente ultrapassa esse intervalo',
    'Repetir pedido: 1 clique gera proposta com mesmos itens da última OV do cliente',
    'Clientes inativos: lista separada de quem não compra há 90+ dias com valor histórico',
    'Análise de crédito: crédito tomado vs limite ajuda a identificar clientes prontos para nova compra',
    'Dashboard mostra top 10 clientes por valor, permitindo priorizar relacionamento',
]
for item in ltv_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('K. O que ainda precisa melhorar', level=2)
improvements = [
    'Tela "Meu Dia" que diga ao vendedor exatamente o que fazer hoje',
    'Follow-up automático quando proposta muda de status',
    'Pipeline ponderado com probabilidade por estágio',
    'Motivo de perda estruturado (dropdown) em vez de texto livre',
    'DRE comercial: resultado real = receita - custo - comissão - impostos',
    'Aging de contas a receber (buckets por dias de atraso)',
    'Migração SQLite → PostgreSQL para suportar mais usuários',
    'Separar app.py em blueprints para manutenção',
    'Testes automatizados para fluxos críticos',
    'Push notifications reais via Web Push API',
    'Integração Omie para sincronizar NF e parcelas',
    'Multi-empresa (ABMT + AEB)',
    'Relatório ad-hoc com filtros combinados livres',
    'Renomear "IA Insights" para "Consulta Rápida" (alinhar expectativa)',
    'Campos obrigatórios no cadastro (UF, segmento, contato)',
    'Concentração de carteira visível (risco de dependência)',
    'Alerta de margem baixa na conversão de proposta',
]
for item in improvements:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('L. Prioridades atuais de desenvolvimento', level=2)
priorities = [
    ('Semana 1-2 (Essencial)', 'Tela "Meu Dia", meta por vendedor, motivo de perda estruturado, campos obrigatórios no cadastro, indicadores de conversão/win rate, expiração automática de propostas'),
    ('Semana 3-4 (Importante)', 'Follow-up automático, pipeline ponderado, DRE comercial, margem consolidada, aging de recebíveis, renomear "IA Insights", alerta de margem na conversão, concentração de carteira, consolidar dashboards'),
    ('Semana 5-8 (Avançado)', 'Migração PostgreSQL, separar blueprints, testes automatizados, push notifications, relatório ad-hoc, registro de concorrente, churn rate'),
    ('Semana 9-12 (Futuro)', 'Integração Omie, multi-empresa, API documentada, relatório semanal por email, categorias configuráveis, analytics de uso interno'),
]
for phase, desc in priorities:
    p = doc.add_paragraph()
    run = p.add_run(f'{phase}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# TABELA MESTRE
# ============================================================
doc.add_heading('TABELA MESTRE — MAPEAMENTO COMPLETO DE MÓDULOS', level=1)

table = doc.add_table(rows=1, cols=10)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Módulo', 'Tela', 'Funcionalidade', 'Status', 'Objetivo Comercial', 'Usuário', 'Dados', 'Regra de Negócio', 'Problema que Resolve', 'Obs']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(7)
            run.font.bold = True
    set_cell_bg(cell, '2D3348')
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

master_data = [
    ['Auth', 'Login', 'Login com sessão', 'Implementado', 'Controle de acesso', 'Todos', 'users', 'Sessão 8h, CSRF, 3 perfis', 'Acesso não autorizado', 'must_change_password no 1o acesso'],
    ['Auth', 'Login', 'Troca de senha obrigatória', 'Implementado', 'Segurança', 'Todos', 'users', 'Flag must_change_password', 'Senha padrão insegura', ''],
    ['Cadastro', 'Cadastros', 'CRUD cliente/fornecedor', 'Implementado', 'Base de dados comercial', 'Todos', 'cadastros', 'CNPJ/CPF único, validação dígito', 'Dados dispersos', 'Campos não obrigatórios'],
    ['Cadastro', 'Cadastro View', 'LTV e saúde do cliente', 'Implementado', 'Retenção de clientes', 'Todos', 'cadastros, OVs', 'Frequência média, dias sem comprar', 'Perda silenciosa de clientes', 'Badges ok/atenção/risco'],
    ['Cadastro', 'Cadastro View', 'Análise de crédito', 'Implementado', 'Gestão de risco', 'Gestor', 'ov_parcelas', 'Tomado/liberado/disponível', 'Vender sem saber se cliente paga', ''],
    ['Cadastro', 'Cadastro View', 'Timeline CRM', 'Implementado', 'Histórico de relacionamento', 'Todos', 'interacoes', '6 tipos de interação', 'Sem registro de contatos', ''],
    ['Proposta', 'Proposta Form', 'Criar proposta venda/compra', 'Implementado', 'Formalizar negociação', 'Todos', 'propostas, itens', 'Numeração auto, 13 categorias', 'Negociação só no WhatsApp', 'Campos específicos por categoria'],
    ['Proposta', 'Proposta View', 'Workflow de status', 'Implementado', 'Funil de vendas', 'Todos', 'propostas, log', 'Transições validadas, 7 status', 'Sem visibilidade de estágio', 'Gestor pode reverter'],
    ['Proposta', 'Proposta View', 'Conversão em OV/OC', 'Implementado', 'Fechar venda/compra', 'Todos', 'propostas→OV/OC', 'Atômico, race condition safe', 'Processo manual de criação de pedido', 'Gera parcelas e comissão auto'],
    ['Proposta', 'Proposta View', 'Duplicar proposta', 'Implementado', 'Agilizar reuso', 'Todos', 'propostas', 'Copia tudo, reseta status', 'Redigitar proposta similar', ''],
    ['Proposta', 'Proposta View', 'Expiração automática', 'Implementado', 'Pipeline limpo', 'Sistema', 'propostas', 'Vencida → Expirada', 'Pipeline inflado com lixo', 'Roda no acesso, não em job'],
    ['OV', 'OV List', 'Listar ordens de venda', 'Implementado', 'Controle de pedidos', 'Todos', 'ordens_venda', 'Vendedor vê só as dele', 'Sem controle de pedidos', ''],
    ['OV', 'OV View', 'Workflow OV expandido', 'Implementado', 'Tracking de status', 'Todos', 'ordens_venda', '6 status incluindo Faturada', 'Sem tracking pós-venda', ''],
    ['OV', 'OV View', 'Parcelas com baixa parcial', 'Implementado', 'Controle financeiro', 'Gestor', 'ov_parcelas', 'Paga/Paga Parcial/Vencida', 'Controle de recebimento manual', 'Audit log em cada baixa'],
    ['OV', 'OV View', 'OCs vinculadas (margem)', 'Implementado', 'Tracking de margem', 'Gestor', 'oc_ov_links', 'Valor alocado compra/venda', 'Sem visibilidade de margem', 'UX básica de vinculação'],
    ['OC', 'OC List/View', 'Ordens de compra', 'Implementado', 'Controle de compras', 'Todos', 'ordens_compra', 'Intermediário, NF, parcelas', 'Compras sem controle', ''],
    ['OC', 'OC View', 'Recebimento parcial', 'Parcial', 'Controle de material', 'Gestor', 'oc_items', 'quantidade_recebida por item', 'Material sem conferência', 'UX limitada'],
    ['Financeiro', 'Parcelas', 'Contas a receber', 'Implementado', 'Fluxo de caixa', 'Gestor', 'ov_parcelas', 'Auto-vencimento, batch', 'Inadimplência invisível', ''],
    ['Financeiro', 'Parcelas', 'Contas a pagar', 'Implementado', 'Fluxo de caixa', 'Gestor', 'oc_parcelas', 'Mesma lógica de OV parcelas', 'Pagamentos sem controle', ''],
    ['Financeiro', 'Dashboard', 'Cashflow 90 dias', 'Implementado', 'Previsão financeira', 'Gestor', 'ov/oc_parcelas', 'Gráfico semanal entradas/saídas', 'Sem visão de caixa futuro', 'Chart.js'],
    ['Dashboard', 'Dashboard', 'KPIs vendas/compras', 'Implementado', 'Visão geral', 'Todos', 'OVs, OCs', 'Filtros combinados', 'Sem visibilidade em tempo real', 'Tabs vendas/compras'],
    ['Dashboard', 'Dashboard', 'Ranking vendedores', 'Implementado', 'Gestão de equipe', 'Gestor', 'OVs, metas', 'Meta vs realizado com %', 'Sem comparação entre vendedores', ''],
    ['Dashboard', 'Dashboard', 'Inadimplência', 'Implementado', 'Cobrança', 'Gestor', 'ov_parcelas', 'Lista vencidas com dias atraso', 'Parcelas vencidas esquecidas', ''],
    ['Dashboard', 'Dashboard', 'Alertas recompra', 'Implementado', 'Retenção/LTV', 'Todos', 'OVs por cadastro', 'Frequência média vs dias sem compra', 'Perda silenciosa de clientes', 'Botão repetir pedido'],
    ['Dashboard', 'Dashboard', 'Previsão faturamento', 'Implementado', 'Planejamento', 'Todos', 'propostas abertas', 'Hoje/semana/pipeline', 'Sem previsão de receita', ''],
    ['Pipeline', 'Pipeline', 'Funil de propostas', 'Implementado', 'Gestão comercial', 'Gestor', 'propostas, OVs', 'Hoje/semana/mês, por vendedor', 'Sem funil de vendas', ''],
    ['Pipeline', 'Pipeline', 'Crédito por cliente', 'Implementado', 'Gestão de risco', 'Gestor', 'ov_parcelas, cadastros', 'Tomado/disponível/status', 'Vender sem saber risco', ''],
    ['Analytics', 'Relatórios', 'Analytics anual', 'Implementado', 'Visão histórica', 'Gestor', 'OVs, OCs', 'Por mês/categoria/vendedor/cliente', 'Sem análise histórica', ''],
    ['Analytics', 'Relatórios', 'Analytics trimestral', 'Implementado', 'Comparativo', 'Gestor', 'OVs, OCs', 'Com comparativo trimestre anterior', 'Sem visão trimestral', ''],
    ['CRM', 'Inteligência', 'Clientes inativos', 'Implementado', 'Reativação', 'Gestor', 'OVs por cadastro', '90+ dias sem compra', 'Clientes perdidos sem detecção', ''],
    ['CRM', 'Assistente', 'Consulta rápida', 'Implementado', 'Acesso rápido a dados', 'Todos', 'OVs, propostas', 'Keyword matching', 'Ter que navegar para achar dado', 'Nome "IA" inadequado'],
    ['Follow-up', 'Follow-ups', 'Tarefas agendadas', 'Implementado', 'Cadência comercial', 'Todos', 'followups', 'Vínculo a entidades, conclusão', 'Esquecer de cobrar retorno', 'Manual, não automático'],
    ['Notas', 'Notas', 'Notas pessoais', 'Implementado', 'Organização pessoal', 'Todos', 'notas', 'Tags, cores, checklist, fixação', 'Informações soltas', ''],
    ['Fechamento', 'Fechamento', 'Comissão mensal', 'Implementado', 'Pagamento justo', 'Gestor', 'OVs, ov_items, config', 'Snapshot JSON, export CSV', 'Comissão calculada na mão', 'Scoring qualitativo'],
    ['Config', 'Configurações', 'ICMS/PIS/Comissão', 'Implementado', 'Flexibilidade fiscal', 'Gestor', 'configuracoes', 'Key-value configurável', 'Mudança fiscal exige deploy', ''],
    ['Busca', 'Topbar', 'Busca global', 'Implementado', 'Acesso rápido', 'Todos', 'busca_global FTS5', 'Prefix search + fallback LIKE', 'Não achar dados rapidamente', ''],
    ['Notificação', 'Topbar/Notif', 'Notificações in-app', 'Implementado', 'Comunicação interna', 'Todos', 'notificacoes', 'Badge + dropdown + lida', 'Sem canal de avisos', 'Sem push real'],
    ['Sugestões', 'Assistente', 'Feedback interno', 'Implementado', 'Melhoria contínua', 'Todos', 'sugestoes', '4 categorias, status de análise', 'Sem canal de feedback', ''],
    ['Metas', 'Dashboard', 'Meta por vendedor', 'Parcial', 'Gestão por objetivos', 'Gestor', 'metas', 'Mensal/semanal por vendedor', 'Sem meta quantificada', 'CRUD incompleto'],
    ['Margem', 'OV View', 'Margem por OV', 'Parcial', 'Rentabilidade', 'Gestor', 'ov_items, oc_ov_links', 'Custo vs venda por item', 'Margem invisível', 'Sem consolidado no dash'],
    ['Meu Dia', '—', 'Dashboard de ação', 'Planejado', 'Adoção diária', 'Vendedor', '—', '—', 'Vendedor não sabe o que fazer', 'Prioridade #1'],
    ['DRE', '—', 'Resultado comercial', 'Planejado', 'Visão de resultado', 'Diretor', '—', '—', 'Sem visão de lucro/prejuízo', ''],
    ['Push', '—', 'Push notifications', 'Planejado', 'Engajamento', 'Todos', '—', '—', 'Notificação só no app aberto', ''],
    ['Omie', '—', 'Integração ERP', 'Planejado', 'Eliminar retrabalho', 'Sistema', '—', '—', 'NF e parcelas em 2 sistemas', ''],
]

for row_data in master_data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(7)
                run.font.name = 'Arial'
        if i == 3:
            color_map = {'Implementado': '1A3A2A', 'Parcial': '2A2A1A', 'Planejado': '1A2A3A', 'Ausente': '3A1A1A'}
            bg = color_map.get(text, None)
            if bg:
                set_cell_bg(cell, bg)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for col_idx in range(10):
    col_widths = [Cm(1.8), Cm(2), Cm(2.5), Cm(1.8), Cm(2.5), Cm(1.5), Cm(2), Cm(3), Cm(3), Cm(2.5)]
    for row in table.rows:
        row.cells[col_idx].width = col_widths[col_idx]

doc.add_page_break()

# ============================================================
# INVENTÁRIO
# ============================================================
doc.add_heading('INVENTÁRIO DO PROJETO', level=1)

doc.add_heading('1. Principais Arquivos', level=2)
files = [
    ('app.py', '~5.300 linhas', 'Backend Flask completo: auth, CRUD, dashboards, analytics, pipeline, fechamento, IA insights, busca, uploads, configs, backup'),
    ('database.py', '~700 linhas', 'Schema SQLite (20+ tabelas), migrations, indexes, seed data, backup, FTS5, get_next_number'),
    ('static/js/app.js', '~5.100 linhas', 'SPA frontend: 60+ ícones SVG, API layer, login, app shell, navegação, dashboard (vendas/compras), OV/OC views, pipeline, analytics, CRM, notas, follow-ups, config, assistente, busca, modals'),
    ('static/js/forms.js', '~800 linhas', 'Formulários: proposta, OV, OC, cadastro. Categorias, subcategorias, UFs, unidades, condições, regimes, ICMS helper'),
    ('static/css/style.css', '~1.470 linhas', 'CSS custom: dark theme, responsive, sidebar, cards, KPIs, badges, tabelas, modais, navegação mobile, animações'),
    ('static/js/sw.js', '43 linhas', 'Service Worker: cache-first para assets, network-only para API, cache CDN'),
    ('templates/index.html', '40 linhas', 'Entry point HTML: carrega CSS/JS, registra SW, detecta atualizações'),
    ('manifest.json', '~20 linhas', 'PWA manifest para instalação no celular'),
]
for name, size, desc in files:
    p = doc.add_paragraph()
    run = p.add_run(f'{name} ({size}): ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2. Principais Componentes', level=2)
components = [
    ('APP object (app.js)', 'Objeto singleton que gerencia toda a SPA: state, navigation, API calls, rendering, theme, search'),
    ('FORMS object (forms.js)', 'Formulários dinâmicos com campos condicionais por categoria de produto'),
    ('api() method', 'Wrapper fetch com CSRF token, timeout 30s, retry, error handling'),
    ('safeAction()', 'Debounce de 1s para prevenir double-click em botões de ação'),
    ('sanitize()', 'Escape de HTML para prevenir XSS'),
    ('LI() helper', 'Função que retorna SVG inline de 60+ ícones Lucide'),
    ('login_required decorator', 'Verifica sessão em todas as rotas de API'),
    ('gestor_required decorator', 'Verifica perfil gerente/diretor para rotas restritas'),
    ('_get_configs()', 'Busca todas as configurações do banco em um dict'),
    ('get_next_number()', 'Gera numeração sequencial para PROP, OV, OC'),
    ('_calc_qualitativo_logic()', 'Score qualitativo de condição de pagamento'),
    ('update_fts()', 'Atualiza índice FTS5 para busca global'),
]
for name, desc in components:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3. Principais Tabelas/Entidades', level=2)
doc.add_paragraph('Ver seção 2.7 para detalhamento completo. Resumo:')
doc.add_paragraph('Domínio principal: users → cadastros → propostas → proposta_items → ordens_venda → ov_items → ov_parcelas (espelhado para compras: ordens_compra → oc_items → oc_parcelas)')
doc.add_paragraph('Suporte: followups, interacoes, notas, anexos, notificacoes, audit_log, busca_global, metas, sugestoes, fechamentos, configuracoes, condicoes_salvas, oc_ov_links, proposta_log')

doc.add_heading('4. Principais Fluxos', level=2)
flows = [
    'Venda: cadastro → proposta VENDA → workflow → conversão OV → workflow OV → parcelas → baixa → fechamento',
    'Compra: cadastro → proposta COMPRA → conversão OC → recebimento → parcelas → baixa',
    'Recompra: alerta sistema → repetir pedido → proposta → ... (mesmo fluxo de venda)',
    'Fechamento: selecionar mês → calcular comissões (vendas + compras) → scoring → fechar → export CSV',
    'Busca: digitar na topbar → FTS5 match → resultados enriquecidos → navegar',
    'Login: username/password → session → CSRF token → renderApp() → navigate(dashboard)',
]
for flow in flows:
    doc.add_paragraph(flow, style='List Bullet')

doc.add_heading('5. Principais Pendências', level=2)
pending = [
    'Tela "Meu Dia" — prioridade #1 para adoção',
    'Follow-up automático por workflow — cadência sem depender de memória',
    'Pipeline ponderado — previsão realista de faturamento',
    'Motivo de perda estruturado — análise "por que perdemos" em escala',
    'DRE Comercial — resultado real por mês',
    'Migração PostgreSQL — estabilidade com múltiplos usuários',
    'Testes automatizados — confiança para evoluir sem regredir',
    'Push notifications — buscar o vendedor proativamente',
]
for item in pending:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6. Funcionalidades Críticas Ausentes', level=2)
critical = [
    ('Tela "Meu Dia"', 'Sem isso, vendedor não tem motivo para abrir o sistema todo dia. É o gatilho de adoção.'),
    ('DRE Comercial', 'Diretor não sabe se a operação está dando lucro ou prejuízo sem exportar dados manualmente.'),
    ('Testes automatizados', 'Comissão e parcelas são cálculos financeiros críticos. Sem testes, qualquer mudança pode gerar erro financeiro silencioso.'),
    ('Migração PostgreSQL', 'SQLite com write lock global não suporta a equipe crescer. É dívida técnica com data de vencimento.'),
    ('Follow-up automático', 'Vendedor esquece de cobrar retorno de propostas enviadas. O sistema deveria criar a tarefa automaticamente.'),
    ('Motivo de perda estruturado', 'Sem dados estruturados de perda, é impossível ajustar estratégia comercial com base em evidência.'),
]
for name, desc in critical:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

# Save
out = 'Documentacao_Sistema_Comercial_ABMT.docx'
doc.save(out)
print(f'OK: {out}')
